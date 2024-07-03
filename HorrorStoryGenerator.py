import sqlalchemy as sa
from pprint import pprint
from langchain_community.utilities import SQLDatabase
from langchain_core.prompts import ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate
from langchain_core.output_parsers import StrOutputParser
from dotenv import load_dotenv
from langchain_groq import ChatGroq
import re
from gradio_client import Client
import os
import shutil
from docx import Document
from PIL import Image
from docx.shared import Inches


class HorrorStoryGenerator:
    def __init__(self, db_uri, model_name, image_model_name, interface_name):
        load_dotenv()
        self.db = SQLDatabase.from_uri(db_uri)
        self.metadata = sa.MetaData()
        self.lingyizhi = sa.Table(
            "lingyizhi",
            self.metadata,
            sa.Column("id", sa.INTEGER, primary_key=True),
            sa.Column("title", sa.TEXT),
            sa.Column("content", sa.TEXT),
        )
        
        if interface_name == "ChatGroq":
            self.chat = ChatGroq(
                temperature=0,
                model=model_name,
                model_kwargs={
                    "frequency_penalty": 0.5,
                    "presence_penalty": 0.5
                },
            )
        else:
            raise ValueError(f"Unsupported interface name: {interface_name}")
        
        self.prompt = ChatPromptTemplate.from_messages([
            SystemMessagePromptTemplate.from_template("你是一名灵异故事创作者。"),
            HumanMessagePromptTemplate.from_template(
"""
请写一篇关于{input}的5000字灵异故事。
输出格式：

- 标题（使用汉语）:
- 汉语关键词（使用汉语 1-3个）:
- 英语关键词（使用英语 1-3个）:
- 内容（使用汉语）:

""")
        ])
        self.chain = self.prompt | self.chat | StrOutputParser()
        self.client = Client(image_model_name)
    
    def query_stories(self, keywords, top_n):
        conditions = [self.lingyizhi.c.title.like(f'%{keyword}%') for keyword in keywords]
        query = sa.select(self.lingyizhi).where(sa.or_(*conditions)).limit(top_n)
        result = self.db.run(query, fetch="cursor")
        stories = list(result.mappings())
        print(f"检索出的故事======\n{stories}")
        return stories
    
    def generate_story(self, input_text):
        content = self.chain.invoke({"input": input_text})
        content = re.sub('\n+', '\n', content)
        print(f"生成的故事======\n{content}")
        return content
    
    def extract_english_keywords(self, story_content):
        pattern = r"汉语关键词：(.+)"
        match = re.search(pattern, story_content)
        if match:
            english_keywords = match.group(1)
            print(f"抽取的关键词======\n{english_keywords}")
            return english_keywords
        else:
            print("没有找到关键词======\n")
            return []
    
    def generate_image(self, prompt, save_directory, save_filename):
        result = self.client.predict(
            prompt=prompt,
            ckpt="4-Step",
            api_name="/generate_image",
        )
        print(f"生成的图像保存路径：{result}")
        if isinstance(result, str) and os.path.isfile(result):
            target_path = os.path.join(save_directory, save_filename)
            os.makedirs(save_directory, exist_ok=True)
            shutil.move(result, target_path)
            print(f"图像已保存到 {target_path}")
            return target_path
        else:
            print("未能正确获取图像文件路径或路径无效。")
            return None
    
    def generate_story_with_image(self, input_text, save_directory, save_filename):
        story_content = self.generate_story(input_text)
        english_keywords = self.extract_english_keywords(story_content)
        if english_keywords:
            image_prompt = ", ".join(english_keywords)
            image_path = self.generate_image(image_prompt, save_directory, save_filename)
        else:
            image_path = None
        return story_content, image_path

    def generate_stories_to_docx(self, docx_filename, keywords, top_n):
        stories = self.query_stories(keywords, top_n)
        document = Document()
        
        for story in stories:
            title = story["title"]
            content, image_path = self.generate_story_with_image(title, "demos/imgs", f"story_{title}.webp")
            document.add_heading(title, level=1)
            
            if image_path:
                png_image_path = 'demos/imgs/converted_image.png'
                with Image.open(image_path) as img:
                    img.save(png_image_path, 'PNG')
                document.add_picture(png_image_path, width=Inches(5))
            document.add_paragraph(content)
        
        doc_save_path = 'demos/docs'
        os.makedirs(doc_save_path, exist_ok=True)
        docx_filename = os.path.join(doc_save_path, docx_filename)
        document.save(docx_filename)


if __name__ == "__main__":
    # 使用示例
    generator = HorrorStoryGenerator(
        db_uri="sqlite:///demos/db/db/horror.db", 
        model_name="llama3-70b-8192", 
        image_model_name="ByteDance/SDXL-Lightning", 
        interface_name="ChatGroq"
    )
    keywords=["公交","地铁"]
    generator.generate_stories_to_docx(
        docx_filename="{keywords}.docx".format(keywords='_'.join(keywords)), 
        keywords=keywords, 
        top_n=10
    )
